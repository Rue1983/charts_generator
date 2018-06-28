import docx
from docx.shared import Inches
from collections import OrderedDict
import gen_charts
import modsec_charts
import modsec_rules
import os
import sqlite3
import configparser
import datetime

RULE_TYPES = 'rule_types.conf'
REPORT_TEMPLATE = 'zywaf防护报告模板2.docx'
DB_NAME = 'alerts0626.db'

dict_insert_pic = {'注：攻击类型详细描述见附录': 'reason_type_pie.png', '从本图中可以获知网站遭遇扫描或者黑客攻击':
    'alerts_by_date.png', '前十大攻击源IP分析': 'ip_source_bar.png', '不可信访问24小时时间分布': 'world1.png',
                   '从24小时的分布图上来看': '24h_stackedline_chart_all.png'}


def rreplace(self, old, new, *max):
    count = len(self)
    if max and str(max[0]).isdigit():
        count = max[0]
    return new.join(self.rsplit(old, count))


def insert_txt_before_para(txt, para):
    """
    :param txt:
    :param para:
    :return:
    """
    new_para = para.insert_paragraph_before()
    new_para.add_run(u'%s' % txt)


def insert_graph_before_para(graph_name, para):
    """
    Use insert_before to avoid indent issue when add pic in new run
    :param graph_name:  picture name
    :param para:  paragraph instance
    :return: none
    """
    new_para = para.insert_paragraph_before()
    run = new_para.add_run(u'')
    run.add_picture('pictures\\%s' % graph_name, width=Inches(6.75))


def get_owasp_types(db_name):
    """
    Get all types of owasp rule id in column owasp
    :return:
    """
    if os.path.isfile(db_name) is False:
        raise FileNotFoundError("Can't find given db %s" % db_name)
    else:
        result = []
        conn = sqlite3.connect(db_name)
        c = conn.cursor()
        cursor = c.execute('SELECT DISTINCT owasp FROM alerts;')
        for row in cursor:
            if row[0] is not None:
                result.append(row[0])
        conn.close()
        print(result)
        return result
#get_owasp_types('alertsbig.db')


def get_majority(input_list):
    majority_ret = []
    tmp = 0
    total = 0
    for r in input_list:
        total += r[1]
    print(total)
    for i, r in enumerate(input_list):
        tmp += r[1]
        if tmp >= total / 2:
            majority_ret = input_list[:i + 1]
            break
    return majority_ret


def get_random_sample(db_name, rule_type):
    """
    Get random sample by given dbname and rule type like: alerts.db, 942
    :param db_name:
    :param rule_type:
    :return: samples in list, total number of the rule types in int.
    """
    if os.path.isfile(db_name) is False:
        raise FileNotFoundError("Can't find given db %s" % db_name)
    else:
        result = []
        times = 0
        conn = sqlite3.connect(db_name)
        c = conn.cursor()
        alerts = c.execute('select ip,time,request,id from alerts where owasp=%s ORDER by random() limit 3' % rule_type)
        for row in alerts:
            #print(row)
            result.append(list(row))
        #print('the result is:', result)
        for sample in result:
            headers = c.execute('select name, value from request_headers where alert_id=%d' % sample[3])
            del (sample[3])
            if not headers:
                continue
            for row in headers:
                sample.append(':'.join(row))
        cursor = c.execute('select count(id) from alerts where owasp=%s' % rule_type)
        for row in cursor:
            times = row[0]
        conn.close()
        #print('the result is:', result, '\n times is : %d' % times)
        return result, times
#get_random_sample(DB_NAME, 942)


def get_config_value(file_name, section, key):
    """
    Get value from given section and keyname from the given conf file.
    :param section:
    :param key:
    :return: value of given section and keyname
    """
    configfile = configparser.ConfigParser()
    configfile.read(RULE_TYPES)
    ret = configfile.get(str(section), key)
    return ret
#get_config_value(RULE_TYPES, '920', 'harm')


def get_config_sections(file_name):
    """
    Get all section names
    :param file_name
    :return: section names in list
    """
    configfile = configparser.ConfigParser()
    configfile.read(RULE_TYPES)
    return configfile.sections()


def gen_customer_report(db_name):
    """
    Generate customer report base on given alerts db and template file.
    :param db_name:
    :return:
    """
    # get summit of 24h and generate 24h trend
    list_24h = [0] * 24
    dict_24h = gen_charts.all_alert_counts_by_reason_24h(db_name)
    for v in dict_24h.values():
        list_24h = list(map(lambda x: x[0] + x[1], zip(list_24h, v)))
    # get upper limit of outlier for daily alert numbers
    upper_limit_24h = gen_charts.get_upper_limit(list_24h)
    print('----------------Upper limit of 24h: ', upper_limit_24h)
    summit_24h = ''
    for i, num in enumerate(list_24h):
        if num > upper_limit_24h:
            summit_24h += '%d点，' % i
    print('+++++++++++++++list 24h is :', list_24h)
    # get deviation days and generate chart by date
    dict_deviation, upper_limit_days = gen_charts.alerts_by_date_chart_pygal(gen_charts.get_alerts_time_reason(db_name))

    # get reason name in Chinese and numbers, count the majority
    top_reasons = gen_charts.get_data_by_reasons(DB_NAME, 'cn')
    reason_majority = get_majority(top_reasons)

    # start to update customer report
    result = 0
    report = docx.Document(REPORT_TEMPLATE)
    for para in report.paragraphs:
        #print(para.text)
        if para.text.startswith('网站日均不可信访问在'):
            summit_days = ''
            multiples = ''
            if dict_deviation:
                for day in dict_deviation.keys():
                    summit_days += '%s年%s月%s日，' % (day[:4], day[4:6].lstrip('0'), day[6:8].lstrip('0'))
                    #print(dict_deviation[day], upper_limit_days, dict_deviation[day]/upper_limit_days)
                    multiples += '%.1f倍，' % (int(dict_deviation[day])/int(upper_limit_days))
                para.text = '网站日均不可信访问在%s次以下属于比较正常范围的。超过%s以上疑似遭到漏洞扫描。' \
                            '%s遭受多轮扫描，不可信流量是离群值上限的%s。' % (upper_limit_days, upper_limit_days,
                                    rreplace(summit_days.rstrip('，'), '，', '和', 1), rreplace(multiples.rstrip('，'),
                                                                                             '，', '和', 1))
        elif para.text.startswith('报告中'):
            analysis = ''
            for r in reason_majority:
                analysis += '“%s”，' % r[0]
            para.text = '报告中%s数量居多，说明' % rreplace(analysis.rstrip('，'), '，', '和', 1)
        elif para.text.startswith('从攻击IP地址分析上看，绝大部分的异常流量来自于'):
            china, foreign, major_area = gen_charts.ip_divide_by_country(DB_NAME)
            analysis = ''
            analysis_china = ''
            analysis_foreign = ''
            city_china = OrderedDict()
            country_foreign = OrderedDict()
            major_china = get_majority(china)  # get majority ip source from the list
            if major_china:
                for i in major_china:
                    city_name, country_name = gen_charts.get_location_by_ip(i[0], 'zh-CN')
                    if city_name in city_china.keys():
                        city_china[city_name] += i[1]
                    else:
                        city_china[city_name] = i[1]
                analysis_china = '分布于%s' % rreplace('，'.join(city_china.keys()), '，', '和', 1)
            major_foreign = get_majority(foreign)  # get majority ip source from the list
            if major_foreign:
                for i in major_foreign:
                    city_name, country_name = gen_charts.get_location_by_ip(i[0], 'zh-CN')
                    if country_name in country_foreign.keys():
                        country_foreign[country_name] += i[1]
                    else:
                        country_foreign[country_name] = i[1]
                analysis_foreign = '分布于%s' % rreplace('，'.join(country_foreign.keys()), '，', '和', 1)
            if major_area == 'China' and analysis_foreign:
                analysis = '从攻击IP地址分析上看，绝大部分的异常流量来自于国内，主要%s。来自国外的流量主要%s。建议将频繁对' \
                           '网站进行不符合安全模型的行为的IP地址配置在IP黑名单中。' % (analysis_china, analysis_foreign)
            elif major_area == 'Foreign' and analysis_china:
                analysis = '从攻击IP地址分析上看，绝大部分的异常流量来自于国外，%s。来自国内的流量主要%s。建议将频繁对' \
                           '网站进行不符合安全模型的行为的IP地址配置在IP黑名单中。' % (analysis_foreign, analysis_china)
            elif major_area == 'China' and not analysis_foreign:
                analysis = '从攻击IP地址分析上看，绝大部分的异常流量来自于国内，主要%s。建议将频繁对' \
                           '网站进行不符合安全模型的行为的IP地址配置在IP黑名单中。' % analysis_foreign
            else:
                raise ValueError("Invalid source region or combination")
            para.text = analysis
        elif para.text.startswith('建议在重要时期重点关注'):
            if summit_24h:
                para.text = '建议在重要时期重点关注%s左右的恶意扫描和黑客攻击行为。' % rreplace(summit_24h.rstrip('，'), '，', '和', 1)
            else:
                para.text = ''
        elif para.text.startswith('注：不可信访问类型详细描述见附录'):
            #print(para.text)
            insert_graph_before_para('reason_type_pie.png', para)
            result += 1
        elif para.text.startswith('前十大攻击源IP分析'):
            #print(para.text)
            insert_graph_before_para('ip_source_bar.png', para)
            result += 1
        elif para.text.startswith('从上图中可以获知网站遭遇扫描或者黑客攻击的高发日期为'):
            summit_days = ''
            if dict_deviation:
                for day in dict_deviation.keys():
                    summit_days += '%s年%s月%s日(%d次)，' % (
                        day[:4], day[4:6].lstrip('0'), day[6:8].lstrip('0'), dict_deviation[day])
                summit_days = '从上图中可以获知网站遭遇扫描或者黑客攻击的高发日期为%s。' % rreplace(summit_days.rstrip('，'), '，', '和', 1)
            insert_graph_before_para('alerts_by_date.png', para)
            para.text = summit_days
            result += 1
        elif para.text.startswith('不可信访问24小时时间分布'):
            insert_graph_before_para('world.png', para)
            result += 1
        elif para.text.startswith('从24小时的分布图上来看'):
            print('=============summit 24h is : ', summit_24h)
            #print(para.text)
            insert_graph_before_para('24h_stackedline_chart_all.png', para)
            if summit_24h:
                analysis = '从24小时的分布图上来看，攻击行为在%s进入高发期。' % summit_24h.rstrip('，')
            else:
                analysis = ''
            para.text = analysis
            result += 1

    # Insert samples in chapter 4.3
    failed_types = get_owasp_types(db_name)
    all_types = get_config_sections(RULE_TYPES)
    all_types = [int(i) for i in all_types]
    no_match_types = list(set(all_types) ^ set(failed_types))
    for ft in all_types:
        for i, t in enumerate(report.tables):
            first_cells = t.rows[0].cells
            if len(first_cells) < 2:
                continue
            if ft in failed_types and first_cells[1].text == get_config_value(RULE_TYPES, ft, 'name'):
                sample, times = get_random_sample(db_name, ft)
                harm = get_config_value(RULE_TYPES, ft, 'harm')
                second_cells = t.rows[1].cells
                sample_text = ''
                for s in sample:
                    sample_text += '\nIP: %s\nTime: %s\nRequest: %s' % (s[0], s[1], s[2])
                    if len(s) > 3:
                        for header in s[3:]:
                            sample_text += '\n%s' % header
                second_cells[1].text = sample_text.lstrip('\n')
                third_cells = t.rows[2].cells
                third_cells[1].text = str(times)
                fourth_cells = t.rows[3].cells
                fourth_cells[1].text = str(harm)
            elif ft in no_match_types and first_cells[1].text == get_config_value(RULE_TYPES, ft, 'name'):
                harm = get_config_value(RULE_TYPES, ft, 'harm')
                fourth_cells = t.rows[3].cells
                fourth_cells[1].text = str(harm)
    report.save("zywaf防护报告模板_test.docx")
    if result != 5:
        raise ValueError("Steps are wrong, result is %d" % result)
#gen_customer_report()


if __name__ == '__main__':
    starttime = datetime.datetime.now()

    modsec_result = modsec_rules.run_rules()
    modsec_charts.owasp_attack_type_waffle(DB_NAME, modsec_result)
    modsec_charts.owasp_attack_type_bar(DB_NAME, modsec_result)
    gen_customer_report(DB_NAME)

    endtime = datetime.datetime.now()
    print('-----total time cost------', endtime - starttime)
    # failed_types = get_owasp_types(DB_NAME)
    # all_types = get_config_sections(RULE_TYPES)
    # all_types = [int(i) for i in all_types]
    # no_match_types = list(set(all_types) ^ set(failed_types))
    # print(failed_types)
    # print('\n', all_types)
    # print('\n', no_match_types)
